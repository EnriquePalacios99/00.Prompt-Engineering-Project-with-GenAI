# app/pages/03_Feedback.py
import streamlit as st
import pandas as pd
from io import StringIO, BytesIO
import json
import re

# Fallback local simple
from services.feedback import summarize_reviews as sum_local, score_sentiment as sent_local

# Gemini (si hay credenciales)
try:
    from services.feedback_gemini import (
        summarize_reviews_gemini as sum_gem,
        score_sentiment_gemini as sent_gem,
        generate_customer_reply_gemini as reply_gem,
    )
    USE_GEMINI = True
except Exception:
    USE_GEMINI = False

st.title("🗣️ Feedback de clientes (Resumen + Sentimiento + Plan + Respuesta)")

# ------------------ Helpers ------------------
def _read(file):
    df = pd.read_csv(file, sep=None, engine="python")
    # Heurística para CSV con ';'
    if df.shape[1] == 1 and ";" in df.columns[0]:
        file.seek(0); df = pd.read_csv(file, sep=";")
    df.columns = [c.strip().lower() for c in df.columns]
    return df

def _slugify(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^a-z0-9\-_.]+", "-", s)
    s = re.sub(r"-{2,}", "-", s).strip("-")
    return s or "reporte-feedback"

def _normalize_ratio_dict(ratio: dict) -> dict:
    """Devuelve fracciones que suman ~1.0; acepta 0–100 o 0–1."""
    vals = {
        "positivo": float(ratio.get("positivo", 0) or 0),
        "neutral":  float(ratio.get("neutral",  0) or 0),
        "negativo": float(ratio.get("negativo", 0) or 0),
    }
    s = sum(vals.values())
    # Si parece porcentaje (p.ej., 100,0,0), conviértelo a fracción
    if s > 1.5:
        vals = {k: v / 100.0 for k, v in vals.items()}
        s = sum(vals.values())
    # Rebalanceo si no suma 1 o si está todo en cero
    if s > 0:
        vals = {k: v / s for k, v in vals.items()}
    else:
        vals = {"positivo": 1/3, "neutral": 1/3, "negativo": 1/3}
    return vals

def _pct(x: float) -> str:
    return f"{x*100:.1f}%"

def _build_summary_csv(summary: dict) -> bytes:
    bullets = summary.get("bullets") or []
    ratio_raw = summary.get("sentiment_ratio", {})
    ratio = _normalize_ratio_dict(ratio_raw)
    plan  = summary.get("action_plan", [])
    row = {
        **{f"bullet_{i+1}": (bullets[i] if i < len(bullets) else "") for i in range(5)},
        "recommendation": summary.get("recommendation", ""),
        "action_plan": " | ".join(plan),
        "customer_reply": summary.get("customer_reply", ""),
        "positivo": ratio.get("positivo", 0.0),
        "neutral": ratio.get("neutral", 0.0),
        "negativo": ratio.get("negativo", 0.0),
        "sample_size": summary.get("sample_size", 0),
    }
    buf = StringIO()
    pd.DataFrame([row]).to_csv(buf, index=False)
    return buf.getvalue().encode("utf-8")

def _build_docx_report(summary: dict, df_sent: pd.DataFrame) -> BytesIO | None:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        st.info("Para exportar a Word instala la librería: `pip install python-docx`")
        return None

    bullets = summary.get("bullets") or []
    ratio = _normalize_ratio_dict(summary.get("sentiment_ratio", {}))
    reco  = summary.get("recommendation", "")
    plan  = summary.get("action_plan", [])
    reply = summary.get("customer_reply", "")
    sample = summary.get("sample_size", 0)

    doc = Document()
    doc.add_heading("Reporte de Feedback de Clientes", level=1)

    p = doc.add_paragraph(f"Muestra analizada: {sample}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("Resumen (3–5 bullets)", level=2)
    for b in bullets:
        doc.add_paragraph(str(b), style="List Bullet")

    doc.add_heading("Recomendación prioritaria", level=2)
    doc.add_paragraph(str(reco))

    doc.add_heading("Plan de acción (3–5 pasos)", level=2)
    for step in plan:
        doc.add_paragraph(str(step), style="List Number")

    doc.add_heading("Respuesta al cliente (plantilla)", level=2)
    doc.add_paragraph(str(reply))

    doc.add_heading("Distribución de sentimiento (aprox.)", level=2)
    doc.add_paragraph(f"Positivo: {_pct(ratio.get('positivo',0.0))}")
    doc.add_paragraph(f"Neutral:  {_pct(ratio.get('neutral',0.0))}")
    doc.add_paragraph(f"Negativo: {_pct(ratio.get('negativo',0.0))}")

    doc.add_heading("Clasificación de sentimientos (muestra)", level=2)
    limit_rows = min(len(df_sent), 1000)
    tbl = doc.add_table(rows=1, cols=3)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Review (≤160c)"
    hdr[1].text = "Sentiment"
    hdr[2].text = "Rationale"

    for _, row in df_sent.head(limit_rows).iterrows():
        r = tbl.add_row().cells
        r[0].text = str(row.get("review",""))
        r[1].text = str(row.get("sentiment",""))
        r[2].text = str(row.get("rationale",""))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
# ---------------------------------------------

# ----------- Carga y selección del CSV -------
up = st.file_uploader("CSV con columna de texto (review/comentario/...)", type=["csv"])

if up:
    df = _read(up)
    st.caption(f"📄 {len(df):,} filas cargadas")
    st.dataframe(df.head(), use_container_width=True)

    cols = df.columns.tolist()
    guess = next((c for c in cols if c in ["review","comentario","texto","opinion","comment"]), cols[0])
    text_col = st.selectbox("Selecciona la columna de texto", cols, index=cols.index(guess))

    # Controles de generación y muestra
    with st.expander("⚙️ Configuración de análisis"):
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            temperature = st.slider("Creatividad (temperature)", 0.0, 1.0, 0.4, 0.1,
                                    help="Más bajo = más determinista para análisis.")
        with c2:
            top_p = st.slider("Top-p", 0.1, 1.0, 0.9, 0.05,
                              help="Probabilidades acumuladas; típicamente 0.9.")
        with c3:
            max_tokens_sum = st.slider("Máx. tokens (resumen+plan+respuesta)", 512, 2048, 768, 64)
        with c4:
            max_tokens_cls = st.slider("Máx. tokens (sentimiento)", 512, 4096, 2048, 64)
        c5, c6 = st.columns(2)
        with c5:
            sample_size_sum = st.number_input("Muestra para resumen", 50, 2000, 300, step=50,
                                              help="Máximo de reviews a considerar.")
        with c6:
            sample_size_cls = st.number_input("Muestra para sentimiento", 50, 2000, 200, step=50,
                                              help="Máximo de reviews a clasificar.")

    reviews_all = df[text_col].dropna().astype(str).tolist()
    reviews_sum = reviews_all[: int(sample_size_sum)]
    reviews_cls = reviews_all[: int(sample_size_cls)]

    # ---------- Un solo botón que corre todo ----------
    if st.button("🧾 Analizar (Resumen + Sentimiento + Plan + Respuesta)", type="primary", use_container_width=True):
        # Resumen (con plan y respuesta)
        try:
            if USE_GEMINI:
                summary = sum_gem(
                    reviews_sum,
                    temperature=temperature,
                    top_p=top_p,
                    max_output_tokens=int(max_tokens_sum),
                    max_reviews=int(sample_size_sum),
                )
            else:
                res_txt = sum_local(reviews_sum)
                summary = {
                    "bullets": [res_txt],
                    "recommendation": "",
                    "sentiment_ratio": {"positivo": 0.34, "neutral": 0.33, "negativo": 0.33},
                    "action_plan": ["Revisar tickets abiertos; CX; 2 semanas"],
                    "customer_reply": "¡Gracias por tu comentario! Escríbenos por DM con tu número de pedido para ayudarte.",
                    "sample_size": len(reviews_sum),
                    "raw": res_txt,
                }
        except Exception as e:
            st.error(f"Error en resumen: {e}")
            summary = {
                "bullets": ["No se pudo generar resumen"],
                "recommendation": "",
                "sentiment_ratio": {"positivo": 0.0, "neutral": 1.0, "negativo": 0.0},
                "action_plan": [],
                "customer_reply": "",
                "sample_size": len(reviews_sum),
            }

        # Sentimiento (por review)
        try:
            if USE_GEMINI:
                rows = sent_gem(
                    reviews_cls,
                    temperature=0.2,
                    top_p=top_p,
                    max_output_tokens=int(max_tokens_cls),
                    max_reviews=int(sample_size_cls),
                )
            else:
                rows = sent_local(reviews_cls)
        except Exception as e:
            st.error(f"Error en sentimiento: {e}")
            rows = sent_local(reviews_cls)

        # ---------- Mostrar resultados ----------
        st.subheader("Resumen (3–5 bullets)")
        bullets = summary.get("bullets") or []
        st.markdown("\n".join([f"- {b}" for b in bullets]))

        st.subheader("Recomendación prioritaria")
        st.write(summary.get("recommendation", ""))

        st.subheader("Plan de acción (3–5 pasos)")
        plan = summary.get("action_plan", [])
        if plan:
            st.markdown("\n".join([f"1. {plan[0]}"] + [f"{i+2}. {p}" for i, p in enumerate(plan[1:])]))
        else:
            st.write("—")

        st.subheader("Respuesta al cliente (plantilla pública)")
        st.text_area("Copia y personaliza si hace falta:", value=summary.get("customer_reply", ""), height=140)

        st.subheader("Distribución de sentimiento (aprox.)")
        ratio = _normalize_ratio_dict(summary.get("sentiment_ratio", {}))
        cma, cmb, cmc = st.columns(3)
        with cma:
            st.metric("Positivo", _pct(ratio["positivo"]))
            st.progress(int(round(ratio["positivo"] * 100)))
        with cmb:
            st.metric("Neutral", _pct(ratio["neutral"]))
            st.progress(int(round(ratio["neutral"] * 100)))
        with cmc:
            st.metric("Negativo", _pct(ratio["negativo"]))
            st.progress(int(round(ratio["negativo"] * 100)))

        st.caption(f"Muestra analizada: {summary.get('sample_size', len(reviews_sum)):,} comentario(s)")

        st.subheader("Análisis de sentimiento (por review)")
        df_sent = pd.DataFrame(rows)
        st.dataframe(df_sent, use_container_width=True)

        # ---------- Descargas ----------
        slug = _slugify(f"feedback-{text_col}")

        # Resumen JSON
        res_json = json.dumps(summary, ensure_ascii=False, indent=2)
        st.download_button(
            "⬇️ Descargar resumen (JSON)",
            res_json.encode("utf-8"),
            file_name=f"{slug}-resumen.json",
            mime="application/json",
            use_container_width=True,
        )

        # Resumen CSV
        st.download_button(
            "⬇️ Descargar resumen (CSV)",
            _build_summary_csv(summary),
            file_name=f"{slug}-resumen.csv",
            mime="text/csv",
            use_container_width=True,
        )

        # Sentimiento CSV
        csv_buf2 = StringIO()
        df_sent.to_csv(csv_buf2, index=False)
        st.download_button(
            "⬇️ Descargar sentimiento (CSV)",
            csv_buf2.getvalue().encode("utf-8"),
            file_name=f"{slug}-sentimiento.csv",
            mime="text/csv",
            use_container_width=True,
        )

        # Word (todo)
        docx_io = _build_docx_report(summary, df_sent)
        if docx_io is not None:
            st.download_button(
                "⬇️ Descargar reporte (Word .docx)",
                data=docx_io,
                file_name=f"{slug}-reporte.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

# ----------- Clasificar y responder un comentario individual -----------
st.markdown("---")
st.subheader("Clasificar y responder un comentario individual")

colx1, colx2 = st.columns([3, 1])
with colx1:
    manual_txt = st.text_area(
        "Pega aquí un comentario/review:",
        height=140,
        placeholder="Ejemplo: 'Llegó tarde y el empaque venía roto. Aún no recibo solución.'",
    )
with colx2:
    brand = st.text_input("Marca (opcional)", value="")

colb1, colb2 = st.columns([1, 3])
with colb1:
    run_single = st.button("🔎 Clasificar y responder")

if run_single and manual_txt.strip():
    try:
        if USE_GEMINI:
            from services.feedback_gemini import score_sentiment_gemini as _cls
            cls = _cls([manual_txt], temperature=0.2, top_p=0.9, max_output_tokens=512, max_reviews=1)
            rep = reply_gem(manual_txt, brand_name=brand or None, temperature=0.4, top_p=0.9, max_output_tokens=512)
        else:
            cls = sent_local([manual_txt])
            rep = {"reply": "Gracias por escribirnos. Queremos ayudarte: envíanos por favor un mensaje con tu número de pedido para revisar el caso."}
    except Exception as e:
        st.error(f"Error en la clasificación o respuesta: {e}")
        cls = [{"review": manual_txt[:160], "sentiment": "neutral", "rationale": ""}]
        rep = {"reply": "Gracias por escribirnos. Queremos ayudarte: envíanos por favor un mensaje con tu número de pedido para revisar el caso."}

    row = cls[0] if cls else {"review": manual_txt[:160], "sentiment": "neutral", "rationale": ""}
    st.write("**Sentimiento:**", row.get("sentiment", "neutral"))
    st.write("**Motivo (rationale):**", row.get("rationale", ""))

    st.write("**Respuesta sugerida:**")
    st.text_area("", value=rep.get("reply",""), height=140)

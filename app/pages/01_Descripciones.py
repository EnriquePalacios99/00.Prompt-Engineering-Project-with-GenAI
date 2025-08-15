import streamlit as st
from io import BytesIO, StringIO
from PIL import Image
from services.llm_gemini import generate_product_description_gemini
import csv, re

st.title("Generación de descripciones (Gemini)")

name = st.text_input("Nombre del producto", "Cereales Ángel")
attrs = st.text_area("Atributos (texto/JSON breve)", "fortificado, crujiente, familiar")

# Mapeo a canales “canónicos” que el prompt entiende mejor
CHANNEL_MAP = {
    "Web (ecommerce)": "Web",
    "Marketplace": "Marketplace",
    "Redes (IG)": "IG",
    "Ads (pagos)": "Ads",
}
channel_label = st.selectbox(
    "Canal",
    list(CHANNEL_MAP.keys()),
    help=(
        "El canal ajusta el tono y la estructura:\n"
        "• IG: cercano/visual, sin tecnicismos.\n"
        "• Ads: breve y directo, foco en beneficio + CTA.\n"
        "• Web: informativo/ordenado, incluye specs si existen.\n"
        "• Marketplace: resalta compatibilidades, medidas, materiales y garantía."
    ),
)
channel = CHANNEL_MAP[channel_label]

# Mensaje corto según canal elegido
CHANNEL_DESC = {
    "IG": "IG: tono cercano y visual (momentos de uso), evita cifras técnicas.",
    "Ads": "Ads: mensajes muy breves y contundentes (beneficio + CTA).",
    "Web": "Web: estilo informativo y ordenado; incluye especificaciones si están en atributos.",
    "Marketplace": "Marketplace: prioriza compatibilidades, tamaño, materiales y garantía si existen.",
}
st.caption(CHANNEL_DESC[channel])

images = st.file_uploader(
    "Imágenes del producto (opcional)",
    type=["jpg","jpeg","png"],
    accept_multiple_files=True,
    help="Adjunta fotos reales del producto para mejorar la precisión (se convierten a JPEG automáticamente)."
)

# Controles con ayuda para no expertos
col1, col2, col3 = st.columns(3)
with col1:
    temperature = st.slider(
        "Creatividad (temperature)",
        min_value=0.0, max_value=1.5, value=0.9, step=0.1,
        help=(
            "Controla lo 'creativo' o aleatorio de la respuesta.\n"
            "• Bajo (0.2–0.5): más factual y consistente.\n"
            "• Medio (0.6–0.8): equilibrio.\n"
            "• Alto (0.9–1.1): más creativo (puede variar más la redacción).\n"
            "Sugerencia para marketing: 0.7–1.0."
        )
    )
with col2:
    top_p = st.slider(
        "Top-p",
        min_value=0.0, max_value=1.0, value=0.95, step=0.05,
        help=(
            "Núcleo de probabilidad: limita las palabras a las más probables hasta acumular 'p'.\n"
            "Más bajo = más estricto; más alto = más variado.\n"
            "Valores típicos: 0.9–0.95."
        )
    )
with col3:
    max_tokens = st.slider(
        "Máx. tokens",
        min_value=256, max_value=2048, value=1024, step=64,
        help=(
            "Longitud máxima de la salida (tokens ≈ trozos de palabra).\n"
            "1 token ≈ 3–4 caracteres aprox.\n"
            "Sube para textos más largos; baja para respuestas más breves."
        )
    )

with st.expander("¿Qué config elijo? (guía rápida y significado)"):
    st.markdown(
        "**¿Qué significa cada parámetro?**\n\n"
        "- **Creatividad (temperature):** determina cuánta variación y riesgo toma el modelo. "
        "Bajo = más literal y consistente; Alto = más creativo y diverso.\n"
        "- **Top-p:** selecciona solo las palabras más probables hasta alcanzar la probabilidad 'p'. "
        "Bajo = más controlado; Alto = más variedad.\n"
        "- **Máx. tokens:** límite superior del tamaño del texto generado. "
        "Si necesitas más detalle, aumenta; si quieres brevedad, reduce.\n\n"
        "**Presets prácticos:**\n"
        "- **Anuncio/Ads muy corto:** temperature **0.7**, top-p **0.9**, **512** tokens → mensajes directos, fuertes y breves.\n"
        "- **Post IG descriptivo:** temperature **0.9**, top-p **0.95**, **768–1024** tokens → tono cercano y visual.\n"
        "- **Ficha Web/Marketplace informativa:** temperature **0.6–0.8**, top-p **0.9**, **1024–1536** tokens "
        "→ texto ordenado con especificaciones cuando existan."
    )

def _to_jpeg_bytes(file):
    """Convierte PNG/JPEG a bytes JPEG (en caso suban PNG, asegura compatibilidad)."""
    data = file.read()
    try:
        img = Image.open(BytesIO(data))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        bio = BytesIO()
        img.save(bio, format="JPEG", quality=92)
        return bio.getvalue()
    except Exception:
        return data

def _coerce_list(value):
    """Asegura una lista de strings limpia desde list/dict/str."""
    if value is None:
        return []
    if isinstance(value, dict):
        value = list(value.values())
    if isinstance(value, str):
        # separa por saltos de línea, viñetas, guiones, punto y coma o coma
        parts = re.split(r"\n|•|-|;|,", value)
        value = [p.strip() for p in parts]
    if isinstance(value, (list, tuple)):
        return [str(x).strip() for x in value if str(x).strip()]
    return [str(value).strip()]

def _slugify(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^a-z0-9\-_.]+", "-", s)
    s = re.sub(r"-{2,}", "-", s).strip("-")
    return s or "descripcion"

def _build_csv_bytes(name, channel, short, long, bullets, hashtags):
    sio = StringIO()
    writer = csv.DictWriter(sio, fieldnames=["name","channel","short","long","bullets","hashtags"])
    writer.writeheader()
    writer.writerow({
        "name": name,
        "channel": channel,
        "short": short,
        "long": long,
        "bullets": " | ".join(bullets),
        "hashtags": " ".join(hashtags),
    })
    return sio.getvalue().encode("utf-8")

def _build_docx_bytes(name, channel, short, long, bullets, hashtags):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        st.info("Para exportar a Word instala la librería: `pip install python-docx`")
        return None

    doc = Document()
    doc.add_heading("Descripción de producto", level=1)
    p = doc.add_paragraph(); r = p.add_run(f"Producto: {name}"); r.font.size = Pt(12)
    p = doc.add_paragraph(); r = p.add_run(f"Canal: {channel}"); r.font.size = Pt(12)

    doc.add_heading("Descripción corta", level=2)
    doc.add_paragraph(short)

    doc.add_heading("Descripción larga (SEO)", level=2)
    doc.add_paragraph(long)

    doc.add_heading("Bullets", level=2)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Hashtags", level=2)
    doc.add_paragraph(" ".join(hashtags))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

if st.button("Generar", type="primary"):
    if not name.strip():
        st.warning("Ingresa el nombre del producto.")
        st.stop()
    if not attrs.strip():
        st.warning("Agrega al menos algunos atributos.")
        st.stop()

    img_bytes = [_to_jpeg_bytes(f) for f in (images or [])]

    out = generate_product_description_gemini(
        name=name,
        attrs_text=attrs,
        channel=channel,
        image_files=img_bytes,
        temperature=temperature,
        top_p=top_p,
        max_tokens=int(max_tokens),
    )

    short = out.get("short", "").strip()
    long_ = out.get("long", "").strip()
    bullets = _coerce_list(out.get("bullets", []))
    hashtags = _coerce_list(out.get("hashtags", []))
    # Solo #hashtags válidos
    hashtags = [h if h.startswith("#") else f"#{h.lstrip('#')}" for h in hashtags]

    st.subheader("Descripción corta")
    st.write(short)

    st.subheader("Descripción larga (SEO)")
    st.write(long_)

    st.subheader("Bullets")
    if bullets:
        st.markdown("\n".join([f"- {b}" for b in bullets]))
    else:
        st.write([])

    st.subheader("Hashtags")
    st.write(" ".join(hashtags) if hashtags else [])

    with st.expander("Salida completa (raw)"):
        st.code(out.get("raw", ""), language="json")

    # === Descargas ===
    slug = _slugify(name)

    # CSV
    csv_bytes = _build_csv_bytes(name, channel, short, long_, bullets, hashtags)
    st.download_button(
        "Descargar CSV",
        data=csv_bytes,
        file_name=f"{slug}.csv",
        mime="text/csv",
        use_container_width=True
    )

    # Word (.docx)
    docx_io = _build_docx_bytes(name, channel, short, long_, bullets, hashtags)
    if docx_io is not None:
        st.download_button(
            "Descargar Word (.docx)",
            data=docx_io,
            file_name=f"{slug}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
